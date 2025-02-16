from docx import Document  # type: ignore
from docx.shared import Pt  # type: ignore
from reportlab.lib.pagesizes import letter  # type: ignore
from reportlab.pdfgen import canvas  # type: ignore
from reportlab.lib.utils import simpleSplit  # type: ignore
from rich.console import Console
import os
import re

console = Console()

def format_text(text):
    """Convert Markdown-style formatting (**bold**, *italic*) to PDF and DOCX compatible formatting."""
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)  # Convert **bold** to <b>...</b>
    text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)  # Convert *italic* to <i>...</i>
    return text

def export_to_pdf(press_kit: str, output_path: str):
    """Exports the press kit to a properly formatted PDF file with bold/italic support."""
    try:
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        max_width = width - 100  # Leave margin space

        y_position = height - 50  # Start from top margin
        line_height = 16  # Adjust spacing

        # Split text into paragraphs
        paragraphs = press_kit.split("\n\n")  

        for paragraph in paragraphs:
            # Convert Rich-style formatting to PDF-friendly HTML-like tags
            formatted_paragraph = format_text(paragraph)
            
            # Wrap text properly within the given width
            wrapped_lines = simpleSplit(formatted_paragraph, "Helvetica", 12, max_width)

            for line in wrapped_lines:
                if "<b>" in line:  # If bold text is detected
                    c.setFont("Helvetica-Bold", 12)
                elif "<i>" in line:  # If italic text is detected
                    c.setFont("Helvetica-Oblique", 12)
                else:
                    c.setFont("Helvetica", 12)

                clean_line = re.sub(r"<\/?[bi]>", "", line)  # Remove tags for drawing
                c.drawString(50, y_position, clean_line)  # Draw text
                y_position -= line_height  # Move down for next line
                
                if y_position < 50:  # Prevent text from going off the page
                    c.showPage()  # Start a new page
                    y_position = height - 50  

            y_position -= line_height  # Extra space after a paragraph

        c.save()
        console.print(f"\n✅ [bold green]Press Kit successfully exported to PDF:[/bold green] {output_path}")

    except Exception as e:
        console.print(f"\n❌ [bold red]Failed to export to PDF:[/bold red] {e}")

def export_to_docx(press_kit: str, output_path: str):
    """Exports the press kit to a formatted DOCX file with bold/italic support."""
    try:
        doc = Document()
        paragraphs = press_kit.split("\n\n")  # Maintain proper paragraph formatting

        for paragraph in paragraphs:
            p = doc.add_paragraph()  # Add a new paragraph

            # Process and format bold & italic text
            matches = re.split(r"(\*\*.*?\*\*|\*.*?\*)", paragraph)

            for match in matches:
                if match.startswith("**") and match.endswith("**"):  # Bold
                    run = p.add_run(match[2:-2])  # Remove **
                    run.bold = True
                elif match.startswith("*") and match.endswith("*"):  # Italic
                    run = p.add_run(match[1:-1])  # Remove *
                    run.italic = True
                else:
                    p.add_run(match)

        doc.save(output_path)
        console.print(f"\n✅ [bold green]Press Kit successfully exported to DOCX:[/bold green] {output_path}")

    except Exception as e:
        console.print(f"\n❌ [bold red]Failed to export to DOCX:[/bold red] {e}")

def export_press_kit(press_kit: str, format_choice: str, company_name: str):
    """Exports the press kit to the selected format."""
    output_folder = "output"
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    if format_choice == "pdf":
        output_path = os.path.join(output_folder, company_name.upper() + ".pdf")
        export_to_pdf(press_kit, output_path)
    elif format_choice == "docx":
        output_path = os.path.join(output_folder, company_name.upper() +".docx")
        export_to_docx(press_kit, output_path)
    else:
        console.print("\n❌ [bold red]Unsupported format choice! Please select either 'pdf' or 'docx'.[/bold red]")

